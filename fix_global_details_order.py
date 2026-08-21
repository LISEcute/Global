from docx import Document

path = r"g_Global\doc\GLOBAL_2026_Modification_Details.docx"
doc = Document(path)

section_texts = {
    "4.8.0 Numerical Refinements",
    "The 4.8.0 numerical changes were proposed by Masahiro Yoshimoto, RIKEN BigRIPS team <masahiro.yoshimoto@riken.jp>.",
    "EDIFF was refined by changing the divisor from 200 to 2000. This makes the energy-change threshold for recalculating charge-changing cross sections about ten times finer.",
    "The integration coefficient sequence was changed from {1,10,100,1000,10000,100000} to {1,1,10,100,1000,10000}. This delays integration-step coarsening by one decade.",
    "These changes are implemented in L_Atima/global/glo_Run.cpp and do not change the ETACHA-style batch input format or the batch CSV output-file naming convention.",
}

for paragraph in list(doc.paragraphs):
    if paragraph.text in section_texts:
        paragraph._element.getparent().remove(paragraph._element)

changed_source = next((p for p in doc.paragraphs if p.text == "Changed Source Files"), None)
if changed_source is None:
    raise RuntimeError("Cannot find Changed Source Files anchor")

details = [
    ("4.8.0 Numerical Refinements", "Heading 1"),
    (
        "The 4.8.0 numerical changes were proposed by Masahiro Yoshimoto, RIKEN BigRIPS team "
        "<masahiro.yoshimoto@riken.jp>.",
        None,
    ),
    (
        "EDIFF was refined by changing the divisor from 200 to 2000. This makes the energy-change "
        "threshold for recalculating charge-changing cross sections about ten times finer.",
        None,
    ),
    (
        "The integration coefficient sequence was changed from {1,10,1000,10000,100000} "
        "to {1,1,10,100,1000,10000}. This delays integration-step coarsening by one decade.",
        None,
    ),
    (
        "These changes are implemented in L_Atima/global/glo_Run.cpp and do not change the "
        "ETACHA-style batch input format or the batch CSV output-file naming convention.",
        None,
    ),
]

for text, style in reversed(details):
    if style:
        changed_source.insert_paragraph_before(text, style=style)
    else:
        changed_source.insert_paragraph_before(text)

doc.save(path)
print(path)
